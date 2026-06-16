from io import BytesIO
from pathlib import Path


class DocumentService:
    TEXT_EXTENSIONS = {"txt", "md"}
    PDF_EXTENSIONS = {"pdf"}
    WORD_EXTENSIONS = {"docx"}

    def extract_text(self, file_content: bytes, filename: str) -> str:
        extension = Path(filename or "").suffix.lower().lstrip(".")
        if extension in self.TEXT_EXTENSIONS:
            return self._extract_plain_text(file_content)
        if extension in self.PDF_EXTENSIONS:
            return self._extract_pdf_text(file_content)
        if extension in self.WORD_EXTENSIONS:
            return self._extract_docx_text(file_content)
        if extension == "doc":
            raise ValueError("暂不支持旧版 .doc 文件，请另存为 .docx 后上传")
        raise ValueError("不支持的文件类型")

    def _extract_plain_text(self, file_content: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return file_content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return file_content.decode("latin-1")

    def _extract_pdf_text(self, file_content: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("缺少 pypdf 依赖，无法读取 PDF 文件") from exc

        reader = PdfReader(BytesIO(file_content))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)

    def _extract_docx_text(self, file_content: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("缺少 python-docx 依赖，无法读取 Word 文件") from exc

        document = Document(BytesIO(file_content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_rows = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_rows.append(" | ".join(cells))
        return "\n".join(paragraphs + table_rows)


document_service = DocumentService()
